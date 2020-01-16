# coding=utf-8
# import pymysql

# connect = pymysql.connect("localhost","root","201619","schat")
# for i in range(1000, 9999):
#     result=int(str(i)[0:2])**2+int(str(i)[2:4])**2
#     result2=(int(str(i)[0:2])+int(str(i)[2:4]))**2
#     if(result2 == i):
#         print(i)
import os
import sys
import stat
from docx import Document
from docx.shared import Inches

document = Document(r'D:\wamp64\www\workspace\schat\wsgi\demo.docx')
document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()
os.chmod(r'D:\workspace', stat.S_IWOTH)
document.save(r'D:\workspace\py\123.docx')
